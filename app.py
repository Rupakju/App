import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import io
import zipfile
import os
import tempfile
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT

# Page configuration
st.set_page_config(
    page_title="Invitation Letter Generator",
    page_icon="📝",
    layout="wide"
)

# Custom CSS for better styling
st.markdown("""
<style>
.main-header {
    text-align: center;
    color: #2c3e50;
    font-size: 2.5rem;
    margin-bottom: 2rem;
}
.sub-header {
    color: #34495e;
    font-size: 1.2rem;
    margin-bottom: 1rem;
}
.success-message {
    background-color: #d4edda;
    color: #155724;
    padding: 1rem;
    border-radius: 0.5rem;
    margin: 1rem 0;
}
</style>
""", unsafe_allow_html=True)

# Title
st.markdown('<h1 class="main-header">📝 Invitation Letter Generator</h1>', unsafe_allow_html=True)
st.markdown('<p class="sub-header">Generate professional invitation letters for visa applications</p>', unsafe_allow_html=True)

# Initialize session state
if 'generated_files' not in st.session_state:
    st.session_state.generated_files = []

def read_word_data(uploaded_file):
    """Read data from uploaded Word file"""
    try:
        doc = Document(uploaded_file)
        info = {}
        
        # Read from first table
        if doc.tables:
            table = doc.tables[0]
            for row in table.rows:
                if len(row.cells) < 2:
                    continue
                key = row.cells[0].text.strip()
                value = row.cells[1].text.strip()
                info[key] = value
        
        return info
    except Exception as e:
        st.error(f"Error reading file: {str(e)}")
        return None

def create_invitation_letter(data, header_image=None, footer_image=None, signature_image=None):
    """Create invitation letter document"""
    try:
        # Create a new Word document
        doc = Document()
        
        # Adjust margins
        section = doc.sections[0]
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.8)
        section.header_distance = Inches(0.2)
        section.footer_distance = Inches(0.1)
        
        # Add header image
        if header_image:
            header = section.header
            header_paragraph = header.paragraphs[0]
            header_paragraph.alignment = 2  # Right-align
            header_run = header_paragraph.add_run()
            header_run.add_picture(header_image, width=Inches(3.44))
        
        # Add footer image
        if footer_image:
            footer = section.footer
            footer_paragraph = footer.paragraphs[0]
            footer_run = footer_paragraph.add_run()
            footer_run.add_picture(footer_image, width=Inches(6.75))
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Lato'
        font.size = Pt(11)
        
        # Add content
        current_date = datetime.now().strftime("%B %d, %Y")
        doc.add_paragraph(f"Date: {current_date}\n")
        
        # To paragraph
        to_para = doc.add_paragraph(f"To\n{data.get('Location of the Bangladesh Embassy that you are applying to (fill address)', 'N/A')}\n")
        
        # Subject line
        subject_para = doc.add_paragraph()
        subject_para.add_run("Subject: Request for business visa to ").bold = False
        subject_para.add_run(f"{data.get('Full Name \n(As it appears on passport)', 'N/A')}, ").bold = True
        subject_para.add_run("Passport No: ").bold = False
        subject_para.add_run(f"{data.get('Passport number', 'N/A')}, ").bold = True
        subject_para.add_run("Nationality: ").bold = False
        subject_para.add_run(f"{data.get('Nationality', 'N/A')}.\n").bold = True
        
        # Body content
        doc.add_paragraph("Dear Sir/Madam,").alignment = 3
        
        doc.add_paragraph("""Save the Children is an international development organization working in 120 countries around the world. The headquarters of Save the Children is located at St Vincent House, 30 Orange Street, London WC2H 7HH, United Kingdom. Save the Children is registered with the NGO Affairs Bureau in Bangladesh (registered number 2630, dated March 20, 2011).""").alignment = 3
        
        # Main body with bold formatting
        body_para = doc.add_paragraph()
        body_para.add_run(f"{data.get('Full Name \n(As it appears on passport)', 'N/A')}, ").bold = True
        body_para.add_run(f"{data.get('Job Title', 'N/A')}, of Save the Children has been invited to the Save the Children International office in Bangladesh to participate in meetings, training, and program activities from ")
        body_para.add_run(f"{data.get('Arrival Date in Bangladesh', 'N/A')} ").bold = True
        body_para.add_run("to ")
        body_para.add_run(f"{data.get('Departure Date', 'N/A')}. ").bold = True
        body_para.add_run("The Save the Children Bangladesh Country Office will ensure all logistical support.").bold = False
        body_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Assistance paragraph
        assistance_para = doc.add_paragraph()
        assistance_para.add_run(f"Your kind assistance in granting a visa for {data.get('Full Name \n(As it appears on passport)', 'N/A')} ")
        assistance_para.add_run("to visit Bangladesh would be highly appreciated. Please contact at Cell no:  +8801913918618 and mail: ")
        assistance_para.add_run("sumon.paul@savethechildren.org ").bold = True
        assistance_para.add_run("if there is any query regarding the processing of this visa.")
        assistance_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        doc.add_paragraph("Thank you for your kind assistance.").alignment = 3
        
        # Add signature image
        if signature_image:
            signature_paragraph = doc.add_paragraph()
            signature_run = signature_paragraph.add_run()
            signature_run.add_picture(signature_image, width=Inches(2.5))
        
        doc.add_paragraph(f"Sumon kumar Paul\nCoordinator - Administration\n")
        
def create_pdf_letter(data, header_image=None, footer_image=None, signature_image=None):
    """Create PDF invitation letter"""
    try:
        # Create PDF buffer
        pdf_buffer = io.BytesIO()
        
        # Create document
        doc = SimpleDocTemplate(
            pdf_buffer,
            pagesize=A4,
            rightMargin=0.8*inch,
            leftMargin=0.9*inch,
            topMargin=1*inch,
            bottomMargin=1*inch
        )
        
        # Get styles
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=12,
            spaceAfter=12,
            alignment=TA_LEFT
        )
        
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=6,
            alignment=TA_LEFT,
            fontName='Helvetica'
        )
        
        justify_style = ParagraphStyle(
            'CustomJustify',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=6,
            alignment=TA_JUSTIFY,
            fontName='Helvetica'
        )
        
        # Story (content list)
        story = []
        
        # Add header image
        if header_image:
            # Save image to temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp:
                tmp.write(header_image.read())
                tmp_path = tmp.name
            
            img = Image(tmp_path, width=3.44*inch, height=1*inch)
            img.hAlign = 'RIGHT'
            story.append(img)
            story.append(Spacer(1, 12))
            
            # Clean up
            os.unlink(tmp_path)
        
        # Add date
        current_date = datetime.now().strftime("%B %d, %Y")
        story.append(Paragraph(f"Date: {current_date}", normal_style))
        story.append(Spacer(1, 12))
        
        # Add To section
        to_text = f"To<br/>{data.get('Location of the Bangladesh Embassy that you are applying to (fill address)', 'N/A')}"
        story.append(Paragraph(to_text, normal_style))
        story.append(Spacer(1, 12))
        
        # Subject line
        subject_text = f"""Subject: Request for business visa to <b>{data.get('Full Name \\n(As it appears on passport)', 'N/A')}</b>, 
        Passport No: <b>{data.get('Passport number', 'N/A')}</b>, 
        Nationality: <b>{data.get('Nationality', 'N/A')}</b>."""
        story.append(Paragraph(subject_text, normal_style))
        story.append(Spacer(1, 12))
        
        # Dear Sir/Madam
        story.append(Paragraph("Dear Sir/Madam,", normal_style))
        story.append(Spacer(1, 12))
        
        # Organization info
        org_text = """Save the Children is an international development organization working in 120 countries around the world. 
        The headquarters of Save the Children is located at St Vincent House, 30 Orange Street, London WC2H 7HH, United Kingdom. 
        Save the Children is registered with the NGO Affairs Bureau in Bangladesh (registered number 2630, dated March 20, 2011)."""
        story.append(Paragraph(org_text, justify_style))
        story.append(Spacer(1, 12))
        
        # Main body
        body_text = f"""<b>{data.get('Full Name \\n(As it appears on passport)', 'N/A')}</b>, 
        {data.get('Job Title', 'N/A')}, of Save the Children has been invited to the Save the Children International office in Bangladesh 
        to participate in meetings, training, and program activities from <b>{data.get('Arrival Date in Bangladesh', 'N/A')}</b> 
        to <b>{data.get('Departure Date', 'N/A')}</b>. 
        The Save the Children Bangladesh Country Office will ensure all logistical support."""
        story.append(Paragraph(body_text, justify_style))
        story.append(Spacer(1, 12))
        
        # Assistance paragraph
        assistance_text = f"""Your kind assistance in granting a visa for {data.get('Full Name \\n(As it appears on passport)', 'N/A')} 
        to visit Bangladesh would be highly appreciated. Please contact at Cell no: +8801913918618 and mail: 
        <b>sumon.paul@savethechildren.org</b> if there is any query regarding the processing of this visa."""
        story.append(Paragraph(assistance_text, justify_style))
        story.append(Spacer(1, 12))
        
        # Thank you
        story.append(Paragraph("Thank you for your kind assistance.", normal_style))
        story.append(Spacer(1, 24))
        
        # Add signature image
        if signature_image:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp:
                tmp.write(signature_image.read())
                tmp_path = tmp.name
            
            sig_img = Image(tmp_path, width=2.5*inch, height=1*inch)
            sig_img.hAlign = 'LEFT'
            story.append(sig_img)
            story.append(Spacer(1, 12))
            
            os.unlink(tmp_path)
        
        # Signature text
        signature_text = "Sumon kumar Paul<br/>Coordinator - Administration"
        story.append(Paragraph(signature_text, normal_style))
        
        # Add footer image
        if footer_image:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp:
                tmp.write(footer_image.read())
                tmp_path = tmp.name
            
            footer_img = Image(tmp_path, width=6.75*inch, height=0.5*inch)
            footer_img.hAlign = 'CENTER'
            story.append(Spacer(1, 24))
            story.append(footer_img)
            
            os.unlink(tmp_path)
        
        # Build PDF
        doc.build(story)
        pdf_buffer.seek(0)
        
        return pdf_buffer.getvalue()
        
    except Exception as e:
        st.error(f"Error creating PDF: {str(e)}")
        return None

# Main interface
col1, col2 = st.columns([1, 1])

with col1:
    st.markdown("### 📁 Upload Files")
    
    # File uploads
    uploaded_files = st.file_uploader(
        "Upload Word files with applicant data",
        type=['docx'],
        accept_multiple_files=True,
        help="Select one or more Word files containing applicant information"
    )
    
    # Image uploads
    st.markdown("### 🖼️ Upload Images (Optional)")
    
    header_image = st.file_uploader(
        "Header Image",
        type=['png', 'jpg', 'jpeg'],
        help="Image for document header"
    )
    
    footer_image = st.file_uploader(
        "Footer Image", 
        type=['png', 'jpg', 'jpeg'],
        help="Image for document footer"
    )
    
    signature_image = st.file_uploader(
        "Signature Image",
        type=['png', 'jpg', 'jpeg'], 
        help="Signature image for the document"
    )

with col2:
    st.markdown("### ⚙️ Generation Options")
    
    # Format selection
    output_format = st.selectbox(
        "Output Format",
        ["Word (.docx)", "PDF", "Both"],
        index=2,
        help="Choose the output format for generated files"
    )
    
    # Generate button
    if st.button("🚀 Generate Invitation Letters", type="primary", use_container_width=True):
        if not uploaded_files:
            st.error("Please upload at least one Word file")
        else:
            # Progress bar
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            generated_files = []
            
            for i, uploaded_file in enumerate(uploaded_files):
                # Update progress
                progress = (i + 1) / len(uploaded_files)
                progress_bar.progress(progress)
                status_text.text(f"Processing {uploaded_file.name}...")
                
                # Read data
                data = read_word_data(uploaded_file)
                if data:
                    base_name = uploaded_file.name.replace('.docx', '')
                    
                    # Generate Word file
                    if output_format in ["Word (.docx)", "Both"]:
                        doc = create_invitation_letter(
                            data, 
                            header_image=header_image, 
                            footer_image=footer_image,
                            signature_image=signature_image
                        )
                        
                        if doc:
                            doc_buffer = io.BytesIO()
                            doc.save(doc_buffer)
                            doc_buffer.seek(0)
                            
                            generated_files.append({
                                'filename': f"{base_name}_invitation.docx",
                                'data': doc_buffer.getvalue(),
                                'type': 'docx',
                                'applicant': data.get('Full Name \n(As it appears on passport)', 'Unknown')
                            })
                    
                    # Generate PDF file
                    if output_format in ["PDF", "Both"]:
                        # Reset image file pointers
                        if header_image:
                            header_image.seek(0)
                        if footer_image:
                            footer_image.seek(0)
                        if signature_image:
                            signature_image.seek(0)
                            
                        pdf_data = create_pdf_letter(
                            data,
                            header_image=header_image,
                            footer_image=footer_image,
                            signature_image=signature_image
                        )
                        
                        if pdf_data:
                            generated_files.append({
                                'filename': f"{base_name}_invitation.pdf",
                                'data': pdf_data,
                                'type': 'pdf',
                                'applicant': data.get('Full Name \n(As it appears on passport)', 'Unknown')
                            })
            
            # Store in session state
            st.session_state.generated_files = generated_files
            
            # Clear progress
            progress_bar.empty()
            status_text.empty()
            
            # Success message
            if generated_files:
                st.success(f"✅ Successfully generated {len(generated_files)} invitation letters!")
            else:
                st.error("❌ No files could be processed")

# Download section
if st.session_state.generated_files:
    st.markdown("### 📥 Download Generated Files")
    
    # Create download options
    col1, col2 = st.columns([1, 1])
    
    with col1:
        # Individual downloads
        st.markdown("**Individual Downloads:**")
        for file_info in st.session_state.generated_files:
            mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document" if file_info['type'] == 'docx' else "application/pdf"
            file_icon = "📄" if file_info['type'] == 'docx' else "📋"
            
            st.download_button(
                label=f"{file_icon} {file_info['applicant']} ({file_info['type'].upper()})",
                data=file_info['data'],
                file_name=file_info['filename'],
                mime=mime_type
            )
    
    with col2:
        # Bulk download as ZIP
        st.markdown("**Bulk Download:**")
        
        # Create ZIP file
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            for file_info in st.session_state.generated_files:
                zip_file.writestr(file_info['filename'], file_info['data'])
        
        zip_buffer.seek(0)
        
        st.download_button(
            label="📦 Download All as ZIP",
            data=zip_buffer.getvalue(),
            file_name=f"invitation_letters_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip",
            mime="application/zip"
        )

# Instructions
with st.expander("📋 Instructions"):
    st.markdown("""
    ### How to use this tool:
    
    1. **Prepare your Word files**: Each file should contain a table with applicant information including:
       - Full Name (As it appears on passport)
       - Passport number
       - Nationality
       - Job Title
       - Arrival Date in Bangladesh
       - Departure Date
       - Location of the Bangladesh Embassy
    
    2. **Upload files**: Select one or more Word files containing applicant data
    
    3. **Upload images** (optional): Add header, footer, and signature images
    
    4. **Generate**: Click the generate button to create invitation letters
    
    5. **Download**: Download individual files or all files as a ZIP archive
    
    ### Supported formats:
    - Input: Word documents (.docx)
    - Images: PNG, JPG, JPEG
    - Output: Word documents (.docx)
    """)

# Footer
st.markdown("---")
st.markdown(
    '<p style="text-align: center; color: #7f8c8d;">Save the Children - Invitation Letter Generator</p>',
    unsafe_allow_html=True
)